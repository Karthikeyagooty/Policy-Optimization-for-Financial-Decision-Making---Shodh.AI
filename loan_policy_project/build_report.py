"""
Build final report (markdown + DOCX) from saved metrics and experiment notes.

This script reads the saved metrics JSON files and produces:
- `reports/final_report.md` (human-readable markdown)
- `reports/final_report.docx` (Microsoft Word document)

Requires: `python-docx` to create the DOCX. Install in venv with `pip install python-docx`.
"""
from pathlib import Path
import json
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).parent
MODEL_DIR = ROOT / 'models'
REPORTS_DIR = ROOT / 'reports'
REPORTS_DIR.mkdir(exist_ok=True)

def load_json(p):
    p = Path(p)
    if p.exists():
        return json.loads(p.read_text())
    return None

grid = load_json(MODEL_DIR / 'grid_best_metrics.json')
final_sup = load_json(MODEL_DIR / 'final_best_metrics.json')
rl = load_json(MODEL_DIR / 'offline_rl_metrics.json')

lines = []
lines.append('# Final Report — Loan Policy Project')
lines.append('')
lines.append('## 1. Overview')
lines.append('This report summarizes the EDA, supervised predictive model, and offline RL experiments.')
lines.append('')
lines.append('## 2. Supervised DL model (MLP)')
if final_sup:
    cfg = final_sup.get('config', {})
    lines.append(f"- Best MLP config: {cfg}")
    lines.append(f"- Test AUC: {final_sup.get('test_auc'):.4f}")
    lines.append(f"- Best Test F1: {final_sup.get('test_f1'):.4f} at threshold {final_sup.get('best_threshold')}")
else:
    lines.append('- No final supervised metrics found.')
lines.append('')
lines.append('## 3. Offline RL agent (Discrete CQL)')
if rl:
    lines.append(f"- Estimated policy value (mean reward per case on test): {rl.get('expected_reward_test', rl.get('expected_reward_test', None))}")
else:
    lines.append('- No offline RL metrics found.')
lines.append('')
lines.append('## 4. Why these metrics?')
lines.append('')
lines.append('**AUC & F1 for DL model:**')
lines.append('- AUC measures the model discrimination across thresholds, useful for ranking risk.')
lines.append('- F1 measures balance between precision and recall at a chosen threshold; important when class imbalance exists and when we care about correct default detection.')
lines.append('')
lines.append('**Estimated Policy Value for RL agent:**')
lines.append('- In offline RL we aim to learn a policy that maximizes expected return. Estimated policy value (average monetary reward on a held-out test set) directly measures the downstream business objective (profit minus losses).')
lines.append('')
lines.append('## 5. Policy comparison and example differences')
lines.append('The supervised DL model defines a thresholding policy (approve when predicted default probability < t). The RL agent learns a policy that maximizes expected reward and may approve some high-risk applicants if the expected reward (interest) outweighs expected losses. Examples with differing decisions can be inspected by comparing model predictions and RL action for the same applicant.')
lines.append('')
lines.append('## 6. Limitations & Future Steps')
lines.append('- Behavior policy was simulated because raw logged approve/deny decisions were unavailable. Replace with real logged propensities if available.')
lines.append('- Use OPE (IS/DR) for robust policy evaluation, and tune CQL/IQL for better performance.')
lines.append('- Consider reward normalization (e.g., divide monetary values by loan amount) and more complex reward that includes time-discounted cashflows.')
lines.append('')
lines.append('## 7. Reproducibility')
lines.append('Follow the instructions in `README.md`. Key scripts:')
lines.append('- `train_mlp_grid.py` — grid search for MLPs')
lines.append('- `evaluate_best.py` — evaluate and save final supervised metrics')
lines.append('- `build_rl_dataset.py` and `train_offline_rl.py` — offline RL pipeline')

md_path = REPORTS_DIR / 'final_report.md'
md_path.write_text('\n'.join(lines))
print('Wrote', md_path)

# Create DOCX
doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

for line in lines:
    if line.startswith('# '):
        doc.add_heading(line[2:], level=1)
    elif line.startswith('## '):
        doc.add_heading(line[3:], level=2)
    else:
        doc.add_paragraph(line)

docx_path = REPORTS_DIR / 'final_report.docx'
doc.save(docx_path)
print('Wrote', docx_path)
